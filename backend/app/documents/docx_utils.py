import io

from docx import Document as DocxDocument


def extract_text(docx_bytes: bytes) -> str:
    """Paragraph + table text, in document order. No pixel bounding boxes are
    available for a DOCX source — the review UI shows a page-level highlight
    for fields grounded from here, not a tight box."""
    doc = DocxDocument(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
